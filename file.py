from mcp.server.fastmcp import FastMCP
from selenium import webdriver
from selenium.webdriver.edge.service import Service as EdgeService
from selenium.webdriver.edge.options import Options as EdgeOptions
from webdriver_manager.microsoft import EdgeChromiumDriverManager
from bs4 import BeautifulSoup
import time
from fpdf import FPDF
from docx import Document
import json

# Inicializar servidor MCP
mcp = FastMCP("WikipediaContentProcessor")

# --- FunciÃ³n para limpiar caracteres especiales ---
def clean_text(text):
    # Elimina caracteres Unicode problemÃ¡ticos y fuera de rango
    return ''.join(char for char in text if ord(char) < 256)

# --- Herramienta: Extraer contenido web de Wikipedia ---
@mcp.tool()
def fetch_wikipedia_content(url: str) -> dict:
    try:
        edge_options = EdgeOptions()
        edge_options.add_argument("--start-maximized")
        driver = webdriver.Edge(service=EdgeService(EdgeChromiumDriverManager().install()), options=edge_options)
        print("ðŸŒ Abriendo navegador...")
        driver.get(url)
        time.sleep(3)  # Esperar a que cargue la pÃ¡gina

        html = driver.page_source
        soup = BeautifulSoup(html, "html.parser")

        print("ðŸ“¥ Extrayendo contenido de Wikipedia...")

        # Buscar el contenedor principal de contenido
        content_div = soup.find("div", {"id": "mw-content-text"})
        if content_div:
            paragraphs = []
            for p in content_div.find_all("p"):  # Encontrar todos los pÃ¡rrafos
                text = p.get_text(strip=True)
                if text:  # Ignorar pÃ¡rrafos vacÃ­os
                    paragraphs.append(clean_text(text))
            return {"status": "ok", "content": paragraphs}
        else:
            return {"status": "error", "message": "âš ï¸ No se encontrÃ³ el contenedor de contenido."}

    except Exception as e:
        return {"status": "error", "message": str(e)}

# --- Herramienta: Mostrar resumen (para TikTok visual) ---
@mcp.tool()
def show_summary(content: list) -> dict:
    summary = ""
    for i, paragraph in enumerate(content[:5], start=1):
        line = f"{i}. {paragraph[:200]}...\n"
        summary += line
    return {"status": "ok", "summary": summary}

# --- Herramienta: Guardar en PDF ---
@mcp.tool()
def save_to_pdf(content: list, filename: str = "contenido.pdf") -> dict:
    try:
        print(f"ðŸ“„ Guardando PDF como '{filename}'...")

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)

        pdf.cell(0, 10, "Contenido Recopilado", ln=True, align='C')
        pdf.ln(10)

        for paragraph in content:
            cleaned_paragraph = clean_text(paragraph)
            pdf.multi_cell(0, 10, cleaned_paragraph)
            pdf.ln(5)

        pdf.output(filename)
        return {"status": "ok", "message": f"âœ… PDF guardado en '{filename}'"}
    except Exception as e:
        return {"status": "error", "message": str(e)}

# --- Herramienta: Guardar en archivo de Word (.docx) ---
@mcp.tool()
def save_to_word(content: list, filename: str = "contenido.docx") -> dict:
    try:
        print(f"ðŸ“˜ Guardando archivo de Word como '{filename}'...")

        doc = Document()
        doc.add_heading('Contenido Recopilado', level=1)

        for paragraph in content:
            doc.add_paragraph(paragraph)

        doc.save(filename)
        return {"status": "ok", "message": f"âœ… Archivo de Word guardado en '{filename}'"}
    except Exception as e:
        return {"status": "error", "message": str(e)}

# --- Prompt principal: Flujo completo ---
@mcp.prompt()
def process_wikipedia_content(url: str) -> str:
    result_fetch = fetch_wikipedia_content(url)
    if result_fetch["status"] != "ok":
        return json.dumps({"error": result_fetch["message"]}, indent=2)

    content = result_fetch["content"]

    result_summary = show_summary(content)
    summary = result_summary.get("summary", "")

    result_pdf = save_to_pdf(content)
    result_word = save_to_word(content)

    return json.dumps({
        "status": "ok",
        "summary": summary,
        "pdf_result": result_pdf,
        "word_result": result_word
    }, indent=2)

# --- Iniciar servidor MCP ---
if __name__ == "__main__":
    print("ðŸš€ Iniciando servidor MCP para procesar contenido de Wikipedia...")
    mcp.run(transport='stdio')
